from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

document = Document()

# Primer parrafo - Centrado con diferentes estilos en el mismo parrafo.
p1 = document.add_paragraph()
p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

run1 = p1.add_run("este es un texto ")
run1.font.name = "Noto Sans"
run1.font.size = Pt(12)

run2 = p1.add_run("con multiples ")
run2.font.name = "Droid Sans"
run2.font.size = Pt(14)
run2.font.bold = True

run3 = p1.add_run("estilos ")
run3.font.name = "Hack"
run3.font.size = Pt(16)
run3.font.italic = True

run4 = p1.add_run("en un mismo parrafo.")
run4.font.name = "Noto Sans"
run4.font.size = Pt(12)
run4.font.underline = True

# segundo párrafo - Justificado con color
p2 = document.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run = p2.add_run("Este es un segundo párrafo con texto justificado y con color azul")
run.font.name = "Hack"
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0, 0, 255)

# Tercer párrafo Alineado a la derecha
p2 = document.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p2.add_run("Este es un tercer párrafo con texto a la derecha y en negrita")
run.font.name = "Hack"
run.font.size = Pt(13)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 0, 255)

# cuarto parrafo alineado a la izquierda con cursiva
p4 = document.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p4.add_run("Este es un cuarto párrafo alineado a la izquierda y en cursiva.")
run.font.name = "DejaVu Serif"
run.font.size = Pt(12)
run.font.italic = True

document.save("documento_multiformato.docx")
