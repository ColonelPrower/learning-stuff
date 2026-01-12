import pandas as pd
from docx import Document
from docx.shared import Pt


class WordAutomation:
    def __init__(self, filename="documento_generado.docx"):
        """inicializa un nuevo documento de word"""
        self.filename = filename
        self.document = Document()

    def agregar_parrafo(self, texto, fuente="Times New Roman", tamano=12):
        """Agrega un parrafo con formato especifico"""
        p = self.document.add_paragraph(texto)
        run = p.runs[0]
        run.font.name = fuente
        run.font.size = Pt(tamano)

    def agregar_tabla_desde_excel(self, excel_file):
        """importa datos de un archivo excel y los inserta en una tabla en el documento"""
        df = pd.read_excel(excel_file)

        table = self.document.add_table(rows=1, cols=len(df.columns))
        hdr_cells = table.rows[0].cells

        # Agregar nombres de columna al encabezado de la tabla
        for i, column in enumerate(df.columns):
            hdr_cells[i].text = column

        # insertar los datos en la tabla
        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, item in enumerate(row):
                row_cells[i].text = str(item)

    def guardar_documento(self):
        """Guarda el documento en el archivo especificado"""
        self.document.save(self.filename)
        print(f"Documento guardado como {self.filename}")
